#!/usr/bin/python3
# *-* coding: utf-8 *-*

import os
import sys
import setproctitle
import uuid
import subprocess
import shlex
# import rbhus_clone_db
import debug
import bcrypt
import argparse
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches
from bs4 import BeautifulSoup
import re

from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtPrintSupport import *

from PyQt5 import QtCore, uic, QtGui, QtWidgets
from PyQt5.QtWidgets import QApplication, QMainWindow, QTreeView, QFileSystemModel, QVBoxLayout, QWidget, QHBoxLayout, QListView
from PyQt5.QtWidgets import QListWidgetItem, QShortcut
from PyQt5.QtGui import QKeySequence, QIcon
from PyQt5.QtMultimedia import QMediaPlayer, QMediaContent
from PyQt5.QtMultimediaWidgets import QVideoWidget
from PyQt5.QtPrintSupport import QPrinter

projDir = os.sep.join(os.path.abspath(__file__).split(os.sep)[:-2])
sys.path.append(projDir)

main_ui_file = os.path.join(projDir,  "tests", "app_test.ui")

# root_folder = "/home/sanath.shetty/Documents/rbhus_clone_root/"

os.environ['QT_LOGGING_RULES'] = "qt5ct.debug=false"

parser = argparse.ArgumentParser(description="Utility to manage assets")
parser.add_argument("-t","--text",dest="text",help="text")
parser.add_argument("-a","--audio",dest="audio",help="audio")
args = parser.parse_args()


class appTest():
    # db = rbhus_clone_db.db()
    def __init__(self):
        # Load main ui
        self.main_ui = uic.loadUi(main_ui_file)
        self.main_ui.setWindowTitle("APP TEST")

        # Text Controls

        self.text_editors = []
        self.create_text_edit()
        
        textLayout = QVBoxLayout()
        textLayout.addWidget(self.text_editors[0])
        self.main_ui.textFrame.setLayout(textLayout)

        if args.text:
            debug.info(args.text)
            self.load_document()

        sizes = ["8", "9", "10", "11", "12", "14", "16", "18", "20", "22", "24", "26", "28", "32", "36", "48", "72"]
        self.main_ui.fontSizeBox.addItems(sizes)
        self.set_font(QFont("Arial"))
        self.main_ui.fontBox.currentFontChanged.connect(self.set_font)
        self.main_ui.fontSizeBox.currentIndexChanged.connect(self.set_font_size)
        
        self.main_ui.boldButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "bold.svg")))
        self.main_ui.italicButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "italic.svg")))
        self.main_ui.underlineButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "underline.svg")))
        self.main_ui.leftAlignButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "align-left.svg")))
        self.main_ui.centerAlignButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "align-center.svg")))
        self.main_ui.rightAlignButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "align-right.svg")))
        self.main_ui.justifyButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "align-justify.svg")))

        self.main_ui.saveButt.setShortcut('Ctrl+S')
        self.main_ui.boldButt.setShortcut(QKeySequence.Bold)
        self.main_ui.italicButt.setShortcut(QKeySequence.Italic)
        self.main_ui.underlineButt.setShortcut(QKeySequence.Underline)

        self.main_ui.saveButt.clicked.connect(self.save_file)
        self.main_ui.printButt.clicked.connect(self.print_file)
        self.main_ui.boldButt.clicked.connect(self.set_bold)
        self.main_ui.italicButt.clicked.connect(self.set_italic)
        self.main_ui.underlineButt.clicked.connect(self.set_underline)
        self.main_ui.leftAlignButt.clicked.connect(self.set_alignment_left)
        self.main_ui.centerAlignButt.clicked.connect(self.set_alignment_center)
        self.main_ui.rightAlignButt.clicked.connect(self.set_alignment_right)
        self.main_ui.justifyButt.clicked.connect(self.set_alignment_justify)

        # Audio Controls

        self.mediaPlayer = QMediaPlayer(None, QMediaPlayer.VideoSurface)
        videoWidget = QVideoWidget()
        if args.audio:
            # fileName = r"C:\Users\Dell\Documents\rbhus_clone_root\template\test_video.mp3"
            fileName = args.audio
        self.mediaPlayer.setMedia(QMediaContent(QUrl.fromLocalFile(fileName)))
        
        self.main_ui.playButton.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "play.svg")))
        self.main_ui.forwardButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "chevron-right.svg")))
        self.main_ui.backwardButt.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "chevron-left.svg")))

        self.main_ui.playButton.clicked.connect(self.play_pause)
        self.main_ui.forwardButt.clicked.connect(self.forward_5_seconds)
        self.main_ui.backwardButt.clicked.connect(self.back_5_seconds)

        self.main_ui.slider.setRange(0, 0)
        self.main_ui.slider.sliderMoved.connect(self.set_position)

        self.main_ui.volumeSlider.setRange(0, 100)
        self.main_ui.volumeSlider.setValue(50)
        self.main_ui.volumeSlider.sliderMoved.connect(self.set_volume)

        videoLayout = QVBoxLayout()
        videoLayout.addWidget(videoWidget)
        self.main_ui.videoFrame.setLayout(videoLayout)
        self.mediaPlayer.setVideoOutput(videoWidget)

        self.mediaPlayer.stateChanged.connect(self.media_state_changed)
        self.mediaPlayer.positionChanged.connect(self.position_changed)
        self.mediaPlayer.durationChanged.connect(self.duration_changed)
        self.mediaPlayer.error.connect(self.handle_error)

        self.launchNudi()

        #Show Window
        self.main_ui.show()
        self.main_ui.update()

        qtRectangle = self.main_ui.frameGeometry()
        centerPoint = QtWidgets.QDesktopWidget().availableGeometry().center()
        qtRectangle.moveCenter(centerPoint)
        self.main_ui.move(qtRectangle.topLeft())

    def launchNudi(self):
        # Specify the path to the application executable
        application_path = r"C:\Program Files (x86)\Nudi 6.1\Nudi 6.1.exe"
        # Use the subprocess module to launch the application in a minimized state
        subprocess.Popen(application_path, creationflags=subprocess.CREATE_NEW_CONSOLE, startupinfo=subprocess.STARTUPINFO(dwFlags=subprocess.STARTF_USESHOWWINDOW))
        # subprocess.run("start "+application_path, shell=True)
 

    def create_text_edit(self):
        text_edit = QTextEdit()
        text_edit.setMinimumSize(595, 842)
        text_edit.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
        # text_edit.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)  # Hide vertical scroll bar
        # text_edit.setFixedHeight(text_edit.fontMetrics().lineSpacing() * self.max_lines)  # Limit height
        text_edit.textChanged.connect(self.handle_text_changed)
        self.text_editors.append(text_edit)        
    
    def load_document(self):
        file_path = args.text
        text_edit = self.text_editors[-1]

        try:
            document = Document(file_path)
            paragraphs = [p.text for p in document.paragraphs]
            text_edit.setText("\n".join(paragraphs))

        except:
            debug.info(str(sys.exc_info()))


    def save_file(self):
        file_path = args.text
        text_edit = self.text_editors[-1]

        try:
            document = Document()
            content = text_edit.toHtml()
            if not content.startswith(' '):
                content = ' ' + content
            document.add_paragraph(content)
            document.save(file_path)

        except:
            debug.info(str(sys.exc_info()))

    def print_file(self):
        dlg = QPrintDialog()
        if dlg.exec_():
            self.text_editors[-1].print_(dlg.printer())

    def handle_text_changed(self):
        current_text_edit = self.text_editors[-1]  # Get the current QTextEdit widget
        document_height = current_text_edit.document().size().height()
        visible_height = current_text_edit.viewport().height()

        if document_height > visible_height:
            self.create_text_edit()
            layout = self.main_ui.textFrame.layout()
            layout.addWidget(self.text_editors[-1])


    def set_font(self, font):
        self.text_editors[-1].setCurrentFont(font)
        self.set_font_size()
        self.set_bold(checked=self.main_ui.boldButt.isChecked())
        self.set_italic(checked=self.main_ui.italicButt.isChecked())
        self.set_underline(checked=self.main_ui.underlineButt.isChecked())

    def set_font_size(self):
        font_size = self.main_ui.fontSizeBox.currentText()
        self.text_editors[-1].setFontPointSize(int(font_size))

    def set_bold(self, checked):
        text_format = self.text_editors[-1].currentCharFormat()
        font = text_format.font()
        font.setBold(checked)
        text_format.setFont(font)
        self.merge_format_on_word_or_selection(text_format)

    def set_italic(self, checked):
        text_format = self.text_editors[-1].currentCharFormat()
        font = text_format.font()
        font.setItalic(checked)
        text_format.setFont(font)
        self.merge_format_on_word_or_selection(text_format)

    def set_underline(self, checked):
        text_format = self.text_editors[-1].currentCharFormat()
        font = text_format.font()
        font.setUnderline(checked)
        text_format.setFont(font)
        self.merge_format_on_word_or_selection(text_format)

    def merge_format_on_word_or_selection(self, format):
        cursor = self.text_editors[-1].textCursor()
        if not cursor.hasSelection():
            cursor.select(QTextCursor.WordUnderCursor)
        cursor.mergeCharFormat(format)
        self.text_editors[-1].mergeCurrentCharFormat

    def set_alignment_left(self):
        self.text_editors[-1].setAlignment(Qt.AlignLeft)

    def set_alignment_center(self):
        self.text_editors[-1].setAlignment(Qt.AlignCenter)

    def set_alignment_right(self):
        self.text_editors[-1].setAlignment(Qt.AlignRight)
    
    def set_alignment_justify(self):
        self.text_editors[-1].setAlignment(Qt.AlignJustify)


    def play_pause(self):
        if self.mediaPlayer.state() == QMediaPlayer.PlayingState:
            self.mediaPlayer.pause()
        else:
            self.mediaPlayer.play()
            
    def media_state_changed(self, state):
        if self.mediaPlayer.state() == QMediaPlayer.PlayingState:
            self.main_ui.playButton.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "pause.svg")))
        else:
            self.main_ui.playButton.setIcon(QtGui.QIcon(os.path.join(projDir, "tests", "image_files", "play.svg")))

    def position_changed(self, position):
        self.main_ui.slider.setValue(position)

    def duration_changed(self, duration):
        self.main_ui.slider.setRange(0, duration)

    def set_position(self, position):
        self.mediaPlayer.setPosition(position)
        # self.mediaPlayer.play()

    def set_volume(self, volume):
        self.mediaPlayer.setVolume(volume)

    def handle_error(self):
        self.main_ui.playButton.setEnabled(False)

    def forward_5_seconds(self):
        self.mediaPlayer.setPosition(self.mediaPlayer.position() + 3000)

    def back_5_seconds(self):
        self.mediaPlayer.setPosition(self.mediaPlayer.position() - 3000)

    def closeEvent(self, event):
        reply = QMessageBox.question(self, 'Exit', 'Are you sure you want to exit?', QMessageBox.Yes | QMessageBox.No,
                                    QMessageBox.No)

        if reply == QMessageBox.Yes:
            event.accept()
        else:
            event.ignore()



if __name__ == '__main__':
    setproctitle.setproctitle("APP_TEST")
    app = QtWidgets.QApplication(sys.argv)
    window = appTest()
    sys.exit(app.exec_())
